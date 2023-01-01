from __future__ import annotations
import os, io, shutil, re, warnings

from typing import List, Dict

from PIL import Image

from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.text.run import Run
from docx.package import Package, OpcPackage

from node import Node, PhotoNode


def convertParagraphsToTree(package: OpcPackage) -> Node:
  """
  Convert a docx file package into internal Node tree structure
  
  Support features:
  1) Each line in Document is converted into a Node
  2) Combine multiple lines into 1 Node with "©©"
  3) Identify photos with "®®"

  Input: OpcPackage, which is the file structure of a docx file\n
  You can get a OpcPackage by following:
  ```python
  f = open('Document.docx', 'rb')
  package = Package.open(f)
  f.close()
  ```
  
  Node parent/children structure is based on Headings in Document

  Example.docx contains:\n
  Heading1\n
  -word1\n
  -Heading2\n
  -- word2

  Node root => Heading1\n
  root.children[1] => word1\n
  root.children[2] => Heading2\n
  root.children[2].children[1] => word2
  """
  # reset image directory first
  shutil.rmtree('image', ignore_errors=True)
  os.mkdir('image')
  
  paragraphs = DocxToNode.getAllParagraphs(package)
  root = Node(None, [])
  curParent = root
  cur_heading_level = 0

  # for loop does not work, https://stackoverflow.com/a/47532461
  i = 0
  while(i < len(paragraphs)):
    p_style = DocxToNode.getParagraphStyle(paragraphs[i]).split()
    
    if p_style[0] != 'heading':
      
      if DocxToNode.isPicture(paragraphs[i]):
        newNode = DocxToNode.createPhotoNote(paragraphs[i], paragraphs[i+1], package, curParent)
        curParent.add(newNode)
        # increment i here 1 more than normal, because a PhotoNode paragraph takes 2 paragraphs
        i += 1
      
      elif DocxToNode.lengthOfBulletList(paragraphs[i]) > 0:
        howManyLinesToSkip = DocxToNode.lengthOfBulletList(paragraphs[i])
        group_paragraphs = paragraphs[i:i+howManyLinesToSkip+1]
        newNode = Node(curParent, group_paragraphs)
        curParent.add(newNode)
        i += howManyLinesToSkip

      # normal paragraph, treat as same level as current level, check if this line is not empty
      elif DocxToNode.isNormalParagraph(paragraphs[i]) and not DocxToNode.isEmptyParagraph(paragraphs[i]):
        newNode = Node(curParent, [paragraphs[i]])
        curParent.add(newNode)

      # If the heading line is actually empty, then skip to next one
      # else:#if DocxToNode.isEmptyParagraph(paragraphs[i]):
    
    elif p_style[0] == 'heading':
    
      # new paragraph has lower(bigger) heading, so move parent node must be higher up, closer to root
      if int(p_style[1]) <= cur_heading_level:
        for _ in range(int(p_style[1]), cur_heading_level + 1):
          curParent = curParent.parent
      
      # This should go in either bigger heading, or smaller heading ( child node ).
      # New node is created under current parent
      new_node = Node(curParent, [paragraphs[i]])
      curParent.add(new_node)
      curParent = new_node
      cur_heading_level = int(p_style[1])
    
    i += 1
  
  return root

class DocxToNode:
  
  @staticmethod
  def getAllParagraphs(docxPackage: OpcPackage) -> List[Paragraph]:
    """
    Convert from python-docx->Package into list of sentences ( python-docx called paragraph )

    Docx documents is stored in XML format.
    Paragraph is basically infos inside <w:p> tags

    ```xml
    <w:p>
      ...
				<w:pStyle w:val="Title"/>
			...
		</w:p>
		<w:p>
			...
				<w:t xml:space="preserve">One sentence will result in </w:t>
			...
				<w:t xml:space="preserve">1</w:t>
			...
				<w:t xml:space="preserve"> Anki note.</w:t>
			...
		</w:p>
    ...
    ```

    You can open a docx document, by:
      1) convert .docx to .zip, and unzip\n
      2) open document.xwl in text editor

    """
    return docxPackage.main_document_part.document.paragraphs
  
  @staticmethod
  def getAllTables(docxPackage: OpcPackage) -> Dict[str, Table]:
    """
    Get all the tables ( 1x1 only ) inside docx, where first line in the table as Key

    Example:
    ```text
    ----------------------
    |¨¨HelloWorldEg      |
    |print("Hello world")|
    ----------------------
    ```
    
    @return  a dictionary { key: `¨¨HelloWorldEg`, value: `¨¨HelloWorldEg  print("Hello world")`
    """
    result = {}
    for t in docxPackage.main_document_part.document.tables:
      result[t.cell(0,0).text.partition('\n')[0]] = t.cell(0,0).text
    return result
  
  @staticmethod
  def getParagraphStyle(para: Paragraph) -> str:
    """
    This return what this sentence ( paragraph )'s hierarchy

    Example:
      title\n
      heading 1\n
      heading 2\n
      normal

    """
    return para.style.name.lower()
  
  @staticmethod
  def getParagraphRuns(para: Paragraph) -> List[Run]:
    """
    Each paragraph has multiple runs, to store different styles

    Example
    ```text
    One sentence will result in **1 Anki** note.
    ```

    This paragraph will have 3 runs:
      - One sentence will result in \n
      - 1 Anki (in bold style )\n
      - note
    """
    return para.runs

  @staticmethod
  def isEmptyParagraph(para: Paragraph) -> bool:
    """
    Check for empty sentence ( paragraph ) in docx file
    """
    return not para.text.replace(' ', '').replace('\t', '').replace('\n', '')
  
  @classmethod
  def isNormalParagraph(cls, para: Paragraph) -> bool:
    """
    Check if this sentence ( paragraph ) is normal, not a heading
    """
    return cls.getParagraphStyle(para).split()[0] == 'normal'

  @staticmethod
  def lengthOfBulletList(para: Paragraph) -> int:
    """
    Check if User want to group multiple paragraphs into 1 Node.

    ©©8 means combine the next 8 lines inside Document into only 1 text node,
     so only 1 Anki Note is created

    Example:
    ```text
    ©©2 List in 1 Anki’s note:
      1) ©©2 means the next 2 lines to be included in 1 note
      2) So this line will be included too
    ```
    All of these will show in 1 Anki note only
    
    """
    if para.text[0:2] == '©©' and para.text[2].isnumeric():
      # return number of lines this list has, indicate after ©©
      return int(para.text.split()[0].replace('©©', ''))
    return -1
  
  @staticmethod
  def isPicture(para: Paragraph) -> bool:
    """
    We define a ®®0 paragraph, the next one is a picture

    There are currently no way to detect a picture in a paragraph automatically.
    So, User must specify ®® on the previous line/paragraph inside the docx document
    
    Also, the digit means do you want to show the picture in 1 Anki note, or multeple

    Rules:
      - ®®0 means treat the picture as 1 note
      - ®®1 means the very next line is a pics, and this pics is going to show on every notes
      created from each line of this heading level in this document
    
    Example:\n
    Heading 1\n
    -texttext1\n
    -®®2\n
    -image1.png\n
    -Heading 2\n
    --texttext2\n

    In this example, 2 Anki Notes is created, texttext1 and texttext2.\n
    Each note has image1.png in it.

    """
    return para.text[0:2] == '®®' and para.text[2].isnumeric()
  
  @staticmethod
  def getImageName(para: Paragraph ) -> str:
    """
    Get the image filename stored inside the docx's xml

    It always start as imageXX.EXT

    Example:
    ```xml
    <w:r>
      <w:rPr/>
      <w:drawing>
      ...
                  <pic:cNvPr id="0" name="image2.png"/>
      ...
      </w:drawing>
    </w:r>
    ```

    @return `**image2.png**`
    """
    if not DocxToNode.getParagraphRuns(para):
      return ""
    cur_xml = DocxToNode.getParagraphRuns(para)[0].element.xml
    regex_match = re.search("image[0-9]*.[a-zA-Z]+", cur_xml)
    if regex_match:
      return regex_match.group(0)
    return ""
  
  @staticmethod
  def getImageIndex(package: OpcPackage, imageName: str) -> int:
    """
    Get the image stored inside document file structure

    `package.image_parts._image_parts[i].partname` get the image path of inner file structure

    /word/media/image1.png\n
    /word/media/image1.png\n
    /word/media/image2.png

    @return the index that match the filename. For image2.png, it returns 2
    """
    document = package.main_document_part.document
    for i in range(len(document.paragraphs)):
      if i >= len(package.image_parts._image_parts):
        raise Exception('The Save function from Microsoft Word create a different image name for each image, please use Google Docs and export as .docx file only')
      
      if imageName in package.image_parts._image_parts[i].partname:
        return i
    return -1
  
  @staticmethod
  def createPhotoNote(paraRR: Paragraph, nextPara: Paragraph, package: OpcPackage, curParent: Node) -> PhotoNode:
    """
    Create a PhotoNode, based on 2 paragraphs
    
    We define a ®®0 paragraph, the next one is a picture.

    Check the docstring on isPicture()

    show_on_children_level: 0 means only show this pic on 1 Anki Note\n
    , 1 means shows on this level's notes\n
    , 2 means this level and next children's level
    """
    imageInfo = [paraRR]
    show_on_children_level = int(paraRR.text[2])

    image_name = DocxToNode.getImageName(nextPara)
    if not image_name:
      warnings.warn("Cannot process image : " + imageInfo[0].text)
      return None

    image_index = DocxToNode.getImageIndex(package, image_name)

    img_binary = package.image_parts._image_parts[image_index].blob
    image = Image.open(io.BytesIO(img_binary))
    image.save('image/'+image_name)

    return PhotoNode(curParent, image_name, image_index, show_on_children_level, imageInfo)




if __name__ == "__main__":
  f = open('Microsoft Word Documents to Anki converter demo.docx', 'rb')
  pp = Package.open(f)
  f.close()

  root = convertParagraphsToTree(pp)
  print(Node.repr(root))
